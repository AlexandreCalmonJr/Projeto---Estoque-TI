import re
import time
from datetime import datetime, timedelta
from io import BytesIO
from pathlib import Path

import pandas as pd
from docx import Document
from flask import abort, current_app, flash, redirect, render_template, request, send_file, session, url_for
from flask_login import current_user, login_required, login_user, logout_user
from werkzeug.utils import secure_filename

from app.models import AssetManager, User, db_execute, db_query, role_required

from . import main_bp


manager = AssetManager()


@main_bp.route('/logo_custom.png')
def serve_custom_logo():
    import os
    from flask import send_from_directory
    from config import basedir
    uploads_dir = os.path.join(basedir, 'uploads')
    logo_path = os.path.join(uploads_dir, 'logo_custom.png')
    if os.path.exists(logo_path):
        return send_from_directory(uploads_dir, 'logo_custom.png')
    # Fallback para o logotipo padrão
    return send_from_directory(os.path.join(current_app.root_path, 'static', 'images'), 'logo.png')


def _allowed_file(filename):
    if not filename or '.' not in filename:
        return False
    extension = filename.rsplit('.', 1)[1].lower()
    return extension in current_app.config['ALLOWED_UPLOAD_EXTENSIONS']


def _normalize_status(status):
    if status == 'Em Manutencao':
        return 'Em Manutenção'
    return status


def _document_templates():
    from app.models import DocumentTemplate
    try:
        tpls = DocumentTemplate.query.all()
        return {t.filename: t.display_name for t in tpls}
    except Exception as e:
        print(f"Erro ao obter templates: {e}")
        return current_app.config['DOCUMENT_TEMPLATES']


def _validate_template_name(template_name):
    if template_name not in _document_templates():
        abort(404)
    from config import basedir
    template_path = Path(basedir) / 'document_templates' / template_name
    templates_dir = template_path.parent.resolve()
    resolved_template = template_path.resolve()
    if templates_dir not in resolved_template.parents or not resolved_template.is_file():
        abort(404)
    return resolved_template


def _read_import_file(file_storage):
    filename = secure_filename(file_storage.filename or '')
    if not _allowed_file(filename):
        raise ValueError('Formato de arquivo inválido. Use CSV ou XLSX.')

    extension = filename.rsplit('.', 1)[1].lower()
    if extension == 'csv':
        try:
            return pd.read_csv(file_storage, dtype=str).fillna('')
        except UnicodeDecodeError:
            file_storage.seek(0)
            return pd.read_csv(file_storage, dtype=str, encoding='latin-1').fillna('')

    return pd.read_excel(file_storage, dtype=str).fillna('')


def get_filter_clauses_and_params():
    categoria_id = request.args.get('categoria', type=int)
    modelo_id = request.args.get('modelo', type=int)
    q = request.args.get('q', '').strip()

    where_clauses = []
    params = {}

    if categoria_id:
        where_clauses.append("a.categoria_id = :categoria_id")
        params['categoria_id'] = categoria_id
    if modelo_id:
        where_clauses.append("a.modelo_id = :modelo_id")
        params['modelo_id'] = modelo_id

    if q:
        where_clauses.append("""
            (a.id_ativo LIKE :search_term OR
             a.numero_serie LIKE :search_term OR
             a.usuario_responsavel LIKE :search_term OR
             a.marca LIKE :search_term OR
             a.localizacao LIKE :search_term OR
             a.destino LIKE :search_term)
        """)
        params['search_term'] = f"%{q}%"

    where_sql = (" WHERE " + " AND ".join(where_clauses)) if where_clauses else ""
    filtros_selecionados = {'categoria_id': categoria_id, 'modelo_id': modelo_id, 'q': q}
    return where_sql, params, filtros_selecionados


@main_bp.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('main.index'))

    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        database_key = request.form.get('database', '').strip()

        user = User.query.filter_by(username=username).first()
        if not user or not user.check_password(password):
            flash('Usuário ou senha inválidos.', 'error')
            return redirect(url_for('main.login'))

        if database_key not in current_app.config['ASSET_DATABASES']:
            flash('Por favor, selecione um banco de dados válido.', 'error')
            return redirect(url_for('main.login'))

        login_user(user)
        session['database_key'] = database_key
        session['database_name'] = current_app.config['ASSET_DATABASES'][database_key]['name']
        return redirect(url_for('main.index'))

    return render_template('login.html', databases=current_app.config['ASSET_DATABASES'])


@main_bp.route('/logout', methods=['POST'])
@login_required
def logout():
    logout_user()
    session.clear()
    flash('Você foi desconectado com segurança.', 'success')
    return redirect(url_for('main.login'))


@main_bp.route('/')
@login_required
def index():
    try:
        where_sql, params, filtros_selecionados = get_filter_clauses_and_params()
        todas_categorias = db_query("SELECT id, nome FROM categorias ORDER BY nome")
        todos_modelos = db_query("SELECT id, nome, categoria_id FROM modelos ORDER BY nome")

        total_ativos = db_query(f"SELECT COUNT(a.id) as count FROM ativos a {where_sql}", params)[0]['count']
        base_join = "FROM ativos a JOIN categorias c ON a.categoria_id = c.id JOIN modelos m ON a.modelo_id = m.id"
        ativos_por_status = db_query(
            f"SELECT a.status, COUNT(a.id) as count {base_join} {where_sql} GROUP BY a.status",
            params,
        )

        destino_query_base = f"SELECT a.destino, COUNT(a.id) as count {base_join} {where_sql}"
        destino_condition = "a.destino IS NOT NULL AND a.destino != ''"
        final_destino_query = f"{destino_query_base} {'AND' if where_sql else 'WHERE'} {destino_condition} GROUP BY a.destino"
        ativos_por_destino = db_query(final_destino_query, params)

        antigos_query_base = f"""
            SELECT a.id_ativo, a.data_aquisicao, c.nome as categoria, m.nome as modelo
            {base_join} {where_sql}
        """
        antigos_condition = "a.status IN ('Em Uso', 'Em Estoque') AND a.data_aquisicao IS NOT NULL AND a.data_aquisicao != ''"
        final_antigos_query = f"{antigos_query_base} {'AND' if where_sql else 'WHERE'} {antigos_condition} ORDER BY a.data_aquisicao ASC LIMIT 5"
        ativos_antigos = db_query(final_antigos_query, params)

        ativos_provisorios_count = db_query("SELECT COUNT(id) as count FROM ativos WHERE numero_serie LIKE 'PROV-%'")[0]['count']

        try:
            almoxarifado_alerta_count = db_query("SELECT COUNT(id) as count FROM consumiveis WHERE quantidade <= estoque_minimo")[0]['count']
            consumiveis_list = db_query("SELECT nome, quantidade, estoque_minimo FROM consumiveis ORDER BY quantidade DESC")
        except Exception:
            almoxarifado_alerta_count = 0
            consumiveis_list = []

        # Warranty alerts
        today_str_idx = datetime.now().strftime('%Y-%m-%d')
        in_60_str_idx = (datetime.now() + timedelta(days=60)).strftime('%Y-%m-%d')
        try:
            garantias_vencendo_count = db_query("""
                SELECT COUNT(id) as count FROM ativos
                WHERE data_garantia IS NOT NULL AND data_garantia != ''
                AND data_garantia >= :hoje AND data_garantia <= :em60
                AND status NOT IN ('Descartado')
            """, {'hoje': today_str_idx, 'em60': in_60_str_idx})[0]['count']
            garantias_vencidas_count = db_query("""
                SELECT COUNT(id) as count FROM ativos
                WHERE data_garantia IS NOT NULL AND data_garantia != ''
                AND data_garantia < :hoje AND status NOT IN ('Descartado')
            """, {'hoje': today_str_idx})[0]['count']
        except Exception:
            garantias_vencendo_count = 0
            garantias_vencidas_count = 0

        lista_ativos_filtrados = db_query(f"""
            SELECT a.id_ativo, c.nome as categoria, m.nome as modelo, a.status, a.usuario_responsavel
            FROM ativos a
            JOIN modelos m ON a.modelo_id = m.id
            JOIN categorias c ON a.categoria_id = c.id
            {where_sql}
            ORDER BY a.id_ativo DESC
        """, params)

        status_labels = [item['status'] for item in ativos_por_status]
        status_data = [item['count'] for item in ativos_por_status]
        destino_labels = [item['destino'] for item in ativos_por_destino]
        destino_data = [item['count'] for item in ativos_por_destino]
        kpis = {
            'total': total_ativos,
            'em_estoque': next((item['count'] for item in ativos_por_status if item['status'] == 'Em Estoque'), 0),
            'em_uso': next((item['count'] for item in ativos_por_status if item['status'] == 'Em Uso'), 0),
            'em_manutencao': next((item['count'] for item in ativos_por_status if item['status'] == 'Em Manutenção'), 0),
        }

        # --- CUSTOS DE MANUTENÇÃO PARA GRÁFICOS ---
        # 1. Evolução Mensal dos Custos (Últimos 6 meses)
        now = datetime.now()
        months_labels = []
        for i in range(5, -1, -1):
            year = now.year
            month = now.month - i
            while month <= 0:
                month += 12
                year -= 1
            months_labels.append(f"{year}-{month:02d}")

        gasto_mensal_res = db_query("""
            SELECT strftime('%Y-%m', data_fim) as mes, SUM(custo) as total
            FROM manutencoes
            WHERE status = 'Concluído'
            GROUP BY mes
        """)
        custos_por_mes = {row['mes']: (row['total'] or 0.0) for row in gasto_mensal_res if row['mes']}
        monthly_costs_data = [float(custos_por_mes.get(m, 0.0)) for m in months_labels]

        # Formatar labels de meses para exibição amigável (Ex: "06/2026")
        months_display_labels = []
        for m in months_labels:
            y, mo = m.split('-')
            months_display_labels.append(f"{mo}/{y}")

        # 2. Custos por Categoria
        custos_categoria_res = db_query("""
            SELECT c.nome as categoria, SUM(m_rec.custo) as total
            FROM manutencoes m_rec
            JOIN ativos a ON m_rec.id_ativo = a.id_ativo
            JOIN categorias c ON a.categoria_id = c.id
            WHERE m_rec.status = 'Concluído'
            GROUP BY c.nome
            ORDER BY total DESC
        """)
        category_cost_labels = [row['categoria'] for row in custos_categoria_res]
        category_cost_data = [float(row['total'] or 0.0) for row in custos_categoria_res]

        # 3. Custos por Marca
        custos_marca_res = db_query("""
            SELECT a.marca, SUM(m_rec.custo) as total
            FROM manutencoes m_rec
            JOIN ativos a ON m_rec.id_ativo = a.id_ativo
            WHERE m_rec.status = 'Concluído'
            GROUP BY a.marca
            ORDER BY total DESC
        """)
        brand_cost_labels = [row['marca'] for row in custos_marca_res]
        brand_cost_data = [float(row['total'] or 0.0) for row in custos_marca_res]

        dashboard_data = {
            'kpis': kpis,
            'status_labels': status_labels,
            'status_data': status_data,
            'destino_labels': destino_labels,
            'destino_data': destino_data,
            'ativos_antigos': ativos_antigos,
            'ativos_provisorios_count': ativos_provisorios_count,
            'almoxarifado_alerta_count': almoxarifado_alerta_count,
            'lista_ativos': lista_ativos_filtrados,
            'consumiveis_labels': [item['nome'] for item in consumiveis_list],
            'consumiveis_quantidades': [item['quantidade'] for item in consumiveis_list],
            'consumiveis_minimos': [item['estoque_minimo'] for item in consumiveis_list],
            'monthly_costs_labels': months_display_labels,
            'monthly_costs_data': monthly_costs_data,
            'category_cost_labels': category_cost_labels,
            'category_cost_data': category_cost_data,
            'brand_cost_labels': brand_cost_labels,
            'brand_cost_data': brand_cost_data,
            'garantias_vencendo_count': garantias_vencendo_count,
            'garantias_vencidas_count': garantias_vencidas_count,
        }
    except Exception as e:
        flash(f"Ocorreu um erro ao carregar os dados do dashboard: {e}", "error")
        dashboard_data = {
            'kpis': {'total': 0, 'em_estoque': 0, 'em_uso': 0, 'em_manutencao': 0},
            'status_labels': [],
            'status_data': [],
            'destino_labels': [],
            'destino_data': [],
            'ativos_antigos': [],
            'ativos_provisorios_count': 0,
            'almoxarifado_alerta_count': 0,
            'lista_ativos': [],
            'consumiveis_labels': [],
            'consumiveis_quantidades': [],
            'consumiveis_minimos': [],
            'monthly_costs_labels': [],
            'monthly_costs_data': [],
            'category_cost_labels': [],
            'category_cost_data': [],
            'brand_cost_labels': [],
            'brand_cost_data': [],
            'garantias_vencendo_count': 0,
            'garantias_vencidas_count': 0,
        }
        todas_categorias, todos_modelos, filtros_selecionados = [], [], {}

    return render_template(
        "index.html",
        title="Dashboard de Ativos",
        data=dashboard_data,
        todas_categorias=todas_categorias,
        todos_modelos=todos_modelos,
        filtros_selecionados=filtros_selecionados,
    )


@main_bp.route('/ativo/novo', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'tecnico')
def form_ativo():
    if request.method == 'POST':
        try:
            manager.registrar_novo_ativo(request.form)
            flash('Ativo cadastrado com sucesso!', 'success')
            return redirect(url_for('main.index'))
        except Exception as e:
            flash(f'Erro ao cadastrar ativo: {e}', 'error')

    categorias = db_query("SELECT *, substr(nome, 1, 2) as sigla FROM categorias ORDER BY nome")
    modelos = db_query("SELECT * FROM modelos ORDER BY nome")
    today = datetime.now().strftime('%Y-%m-%d')
    return render_template(
        "form_ativo.html",
        title="Cadastramento de Ativos de TI",
        categorias=categorias,
        modelos=modelos,
        today=today,
    )


@main_bp.route('/ativo/<id_ativo>')
@login_required
def detalhes_ativo(id_ativo):
    query_ativo = """
        SELECT a.*, m.nome as modelo, c.nome as categoria
        FROM ativos a
        JOIN modelos m ON a.modelo_id = m.id
        JOIN categorias c ON a.categoria_id = c.id
        WHERE a.id_ativo = :id_ativo
    """
    ativo_list = db_query(query_ativo, {'id_ativo': id_ativo})
    if not ativo_list:
        flash('Ativo não encontrado.', 'error')
        return redirect(url_for('main.index'))

    historico = db_query(
        "SELECT * FROM historico WHERE id_ativo = :id_ativo ORDER BY timestamp DESC",
        {'id_ativo': id_ativo},
    )
    termo = db_query(
        "SELECT * FROM termos WHERE id_ativo = :id_ativo ORDER BY data_criacao DESC LIMIT 1",
        {'id_ativo': id_ativo},
    )
    termo = termo[0] if termo else None

    manutencoes = db_query(
        "SELECT * FROM manutencoes WHERE id_ativo = :id_ativo ORDER BY data_inicio DESC",
        {'id_ativo': id_ativo}
    )

    # Build responsibility history from historico events
    responsaveis_historico = []
    for h in historico:
        if h['evento'] in ('Movimentação', 'Distribuição') and h['detalhes'] and 'responsável' in h['detalhes'].lower():
            import re as _re
            match = _re.search(r'Novo responsável: (.+?)[\.\s]*(?:Chamado:|$)', h['detalhes'], _re.IGNORECASE)
            chamado_match = _re.search(r'Chamado: ([\w\-/]+)', h['detalhes'], _re.IGNORECASE)
            if match:
                responsaveis_historico.append({
                    'responsavel': match.group(1).strip().rstrip('.'),
                    'data': h['timestamp'],
                    'chamado': chamado_match.group(1) if chamado_match else None,
                })

    return render_template(
        "detalhes_ativo.html",
        title=f"Detalhes: {ativo_list[0]['id_ativo']}",
        ativo=ativo_list[0],
        historico=historico,
        termo=termo,
        manutencoes=manutencoes,
        document_templates=_document_templates(),
        responsaveis_historico=responsaveis_historico,
        now=datetime.now,
        timedelta=timedelta,
    )


@main_bp.route('/exportar-excel')
@login_required
def exportar_excel():
    """Exports current filtered assets list as a formatted .xlsx file."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    where_sql, params, _ = get_filter_clauses_and_params()
    ativos = db_query(f"""
        SELECT a.id_ativo, c.nome as categoria, m.nome as modelo, a.marca,
               a.numero_serie, a.status, a.usuario_responsavel, a.localizacao,
               a.data_aquisicao, a.data_garantia, a.fornecedor
        FROM ativos a
        JOIN modelos m ON a.modelo_id = m.id
        JOIN categorias c ON a.categoria_id = c.id
        {where_sql}
        ORDER BY a.id_ativo
    """, params)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Ativos"

    # Header style
    header_fill = PatternFill("solid", fgColor="1E3A5F")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    headers = [
        "Patrimônio", "Categoria", "Modelo", "Marca",
        "Nº Série", "Status", "Responsável", "Localização",
        "Data Aquisição", "Venc. Garantia", "Fornecedor"
    ]

    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = thin_border

    ws.row_dimensions[1].height = 22

    # Data rows with zebra striping
    alt_fill = PatternFill("solid", fgColor="EEF2FF")
    for row_idx, ativo in enumerate(ativos, start=2):
        row_data = [
            ativo['id_ativo'], ativo['categoria'], ativo['modelo'], ativo['marca'],
            ativo['numero_serie'], ativo['status'], ativo['usuario_responsavel'] or '',
            ativo['localizacao'] or '', ativo['data_aquisicao'] or '',
            ativo['data_garantia'] or '', ativo['fornecedor'] or ''
        ]
        fill = alt_fill if row_idx % 2 == 0 else None
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(vertical='center')
            if fill:
                cell.fill = fill

    # Auto-fit columns
    col_widths = [15, 15, 20, 15, 20, 15, 20, 20, 15, 15, 20]
    for i, width in enumerate(col_widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = width

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    filename = f"ativos_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename
    )


@main_bp.route('/relatorio/colaborador')
@login_required
def relatorio_colaborador():
    """Report of all assets assigned to a given employee."""
    q = request.args.get('q', '').strip()
    ativos = []
    totais_por_categoria = []
    if q:
        ativos = db_query("""
            SELECT a.id_ativo, c.nome as categoria, m.nome as modelo,
                   a.marca, a.status, a.localizacao, a.usuario_responsavel,
                   a.data_aquisicao, a.data_garantia
            FROM ativos a
            JOIN modelos m ON a.modelo_id = m.id
            JOIN categorias c ON a.categoria_id = c.id
            WHERE LOWER(a.usuario_responsavel) LIKE LOWER(:q)
            ORDER BY c.nome, a.id_ativo
        """, {'q': f'%{q}%'})

        # Group by category for summary
        cat_count = {}
        for a in ativos:
            cat = a['categoria']
            cat_count[cat] = cat_count.get(cat, 0) + 1
        totais_por_categoria = [{'categoria': k, 'total': v} for k, v in sorted(cat_count.items())]

    return render_template(
        'relatorio_colaborador.html',
        title='Ativos por Colaborador',
        ativos=ativos,
        totais_por_categoria=totais_por_categoria,
        q=q,
        now=datetime.now,
        timedelta=timedelta,
    )


@main_bp.route('/notificacoes')
@login_required
def notificacoes():
    """Returns JSON with internal notifications for the bell icon."""
    from flask import jsonify
    alertas = []
    today_str = datetime.now().strftime('%Y-%m-%d')
    in_60_days = (datetime.now() + timedelta(days=60)).strftime('%Y-%m-%d')
    in_30_days = (datetime.now() + timedelta(days=30)).strftime('%Y-%m-%d')

    try:
        # Assets in maintenance for more than 15 days
        manut_antigas = db_query("""
            SELECT COUNT(*) as cnt FROM manutencoes m
            JOIN ativos a ON m.id_ativo = a.id_ativo
            WHERE a.status = 'Em Manutenção'
            AND m.status = 'Em Manutenção'
            AND julianday('now') - julianday(m.data_inicio) > 15
        """)
        cnt_manut = manut_antigas[0]['cnt'] if manut_antigas else 0
        if cnt_manut > 0:
            alertas.append({'tipo': 'warning', 'icone': '🔧',
                            'mensagem': f'{cnt_manut} ativo(s) em manutenção há mais de 15 dias',
                            'link': '/manutencao'})

        # Warranties expiring in 30 days
        garantias_vencendo = db_query("""
            SELECT COUNT(*) as cnt FROM ativos
            WHERE data_garantia IS NOT NULL
            AND data_garantia != ''
            AND data_garantia >= :hoje
            AND data_garantia <= :em_30
        """, {'hoje': today_str, 'em_30': in_30_days})
        cnt_gar30 = garantias_vencendo[0]['cnt'] if garantias_vencendo else 0
        if cnt_gar30 > 0:
            alertas.append({'tipo': 'warning', 'icone': '⚠️',
                            'mensagem': f'{cnt_gar30} garantia(s) vencem em 30 dias',
                            'link': '/'})

        # Warranties already expired
        garantias_vencidas = db_query("""
            SELECT COUNT(*) as cnt FROM ativos
            WHERE data_garantia IS NOT NULL
            AND data_garantia != ''
            AND data_garantia < :hoje
            AND status NOT IN ('Descartado')
        """, {'hoje': today_str})
        cnt_venc = garantias_vencidas[0]['cnt'] if garantias_vencidas else 0
        if cnt_venc > 0:
            alertas.append({'tipo': 'danger', 'icone': '🔴',
                            'mensagem': f'{cnt_venc} ativo(s) com garantia vencida',
                            'link': '/'})

        # Consumables below minimum
        consumiveis_criticos = db_query("""
            SELECT COUNT(*) as cnt FROM consumiveis WHERE quantidade <= estoque_minimo
        """)
        cnt_cons = consumiveis_criticos[0]['cnt'] if consumiveis_criticos else 0
        if cnt_cons > 0:
            alertas.append({'tipo': 'danger', 'icone': '📦',
                            'mensagem': f'{cnt_cons} consumivel(is) abaixo do estoque mínimo',
                            'link': '/consumiveis'})
    except Exception as e:
        alertas.append({'tipo': 'info', 'icone': 'ℹ️', 'mensagem': f'Erro ao carregar alertas: {e}', 'link': '#'})

    return jsonify({'alertas': alertas, 'total': len(alertas)})


@main_bp.route('/manutencao')
@login_required
def manutencao():
    # 1. Ativos em Manutenção
    ativos_manutencao = db_query("""
        SELECT a.id_ativo, a.numero_serie, a.marca, m.nome as modelo, c.nome as categoria,
               m_rec.data_inicio, m_rec.numero_chamado, m_rec.id as manutencao_id
        FROM ativos a
        JOIN modelos m ON a.modelo_id = m.id
        JOIN categorias c ON a.categoria_id = c.id
        LEFT JOIN (
            SELECT id, id_ativo, data_inicio, numero_chamado
            FROM manutencoes
            WHERE status = 'Em Manutenção'
        ) m_rec ON a.id_ativo = m_rec.id_ativo
        WHERE a.status = 'Em Manutenção'
        ORDER BY m_rec.data_inicio DESC
    """)

    # 2. Histórico de Manutenções (Concluídas)
    historico_manutencoes = db_query("""
        SELECT m_rec.id as manutencao_id, m_rec.id_ativo, m_rec.numero_chamado, m_rec.data_inicio, m_rec.data_fim,
               m_rec.custo, m_rec.pecas_substituidas, m_rec.descricao,
               a.marca, m.nome as modelo, c.nome as categoria
        FROM manutencoes m_rec
        JOIN ativos a ON m_rec.id_ativo = a.id_ativo
        JOIN modelos m ON a.modelo_id = m.id
        JOIN categorias c ON a.categoria_id = c.id
        WHERE m_rec.status = 'Concluído'
        ORDER BY m_rec.data_fim DESC
    """)

    # 3. KPIs Financeiros e Quantidades
    current_month = datetime.now().strftime('%Y-%m')
    
    # Total gasto no mês
    gasto_mes_res = db_query("""
        SELECT SUM(custo) as total FROM manutencoes
        WHERE status = 'Concluído' AND strftime('%Y-%m', data_fim) = :current_month
    """, {'current_month': current_month})
    gasto_mes = gasto_mes_res[0]['total'] if gasto_mes_res and gasto_mes_res[0]['total'] is not None else 0.0

    # Total gasto geral
    gasto_geral_res = db_query("""
        SELECT SUM(custo) as total FROM manutencoes
        WHERE status = 'Concluído'
    """)
    gasto_geral = gasto_geral_res[0]['total'] if gasto_geral_res and gasto_geral_res[0]['total'] is not None else 0.0

    # Ativos em reparo
    ativos_em_reparo = len(ativos_manutencao)

    # Marca/Modelo crítico
    critico_res = db_query("""
        SELECT a.marca, m.nome as modelo, SUM(m_rec.custo) as total_custo
        FROM manutencoes m_rec
        JOIN ativos a ON m_rec.id_ativo = a.id_ativo
        JOIN modelos m ON a.modelo_id = m.id
        WHERE m_rec.status = 'Concluído'
        GROUP BY a.marca, m.nome
        ORDER BY total_custo DESC LIMIT 1
    """)
    critico = f"{critico_res[0]['marca']} {critico_res[0]['modelo']}" if critico_res else "Nenhum"

    kpis = {
        'gasto_mes': gasto_mes,
        'gasto_geral': gasto_geral,
        'ativos_em_reparo': ativos_em_reparo,
        'critico': critico
    }

    return render_template(
        "manutencao.html",
        title="Controle de Manutenção & Custos",
        ativos=ativos_manutencao,
        historico=historico_manutencoes,
        kpis=kpis
    )


@main_bp.route('/ativo/<id_ativo>/concluir_manutencao', methods=['POST'])
@login_required
@role_required('admin', 'tecnico')
def concluir_manutencao(id_ativo):
    try:
        custo_raw = request.form.get('custo', '0').strip().replace(',', '.')
        try:
            custo = float(custo_raw)
        except ValueError:
            custo = 0.0

        pecas = request.form.get('pecas_substituidas', '').strip()
        descricao = request.form.get('descricao', '').strip()
        novo_status = _normalize_status(request.form.get('novo_status', 'Em Estoque').strip())

        if novo_status not in ['Em Estoque', 'Em Uso']:
            novo_status = 'Em Estoque'

        # Busca registro ativo de manutenção
        manutencao_list = db_query("""
            SELECT id, numero_chamado FROM manutencoes
            WHERE id_ativo = :id_ativo AND status = 'Em Manutenção'
            ORDER BY id DESC LIMIT 1
        """, {'id_ativo': id_ativo})

        if manutencao_list:
            manut_id = manutencao_list[0]['id']
            chamado = manutencao_list[0]['numero_chamado']
            
            db_execute("""
                UPDATE manutencoes
                SET data_fim = :data_fim, custo = :custo, pecas_substituidas = :pecas, descricao = :descricao, status = 'Concluído'
                WHERE id = :id
            """, {
                'data_fim': datetime.now().strftime('%Y-%m-%d'),
                'custo': custo,
                'pecas': pecas,
                'descricao': descricao,
                'id': manut_id
            })
        else:
            # Caso legado: cria e conclui na hora
            chamado = request.form.get('numero_chamado', 'LEGADO').strip() or 'LEGADO'
            db_execute("""
                INSERT INTO manutencoes (id_ativo, numero_chamado, data_inicio, data_fim, custo, pecas_substituidas, descricao, status)
                VALUES (:id_ativo, :chamado, :hoje, :hoje, :custo, :pecas, :descricao, 'Concluído')
            """, {
                'id_ativo': id_ativo,
                'chamado': chamado,
                'hoje': datetime.now().strftime('%Y-%m-%d'),
                'custo': custo,
                'pecas': pecas,
                'descricao': descricao
            })

        # Altera o status do ativo
        manager.movimentar(id_ativo, novo_status, chamado)

        # Log do histórico geral
        log_detalhes = f"Manutenção concluída (Chamado: {chamado}). Custo: R$ {custo:.2f}. Peças: {pecas or 'Nenhuma'}. Descrição: {descricao or 'Nenhuma'}."
        db_execute("""
            INSERT INTO historico (id_ativo, evento, detalhes)
            VALUES (:id_ativo, 'Movimentação', :detalhes)
        """, {'id_ativo': id_ativo, 'detalhes': log_detalhes})

        flash('Manutenção concluída com sucesso!', 'success')
        return redirect(url_for('main.detalhes_ativo', id_ativo=id_ativo))
    except Exception as e:
        flash(f"Erro ao concluir manutenção: {e}", 'error')
        return redirect(url_for('main.detalhes_ativo', id_ativo=id_ativo))


@main_bp.route('/manutencao/excluir/<int:id>', methods=['POST'])
@login_required
def excluir_manutencao(id):
    try:
        if not current_user.is_admin:
            flash('Apenas administradores podem excluir registros de manutenção.', 'error')
            return redirect(url_for('main.manutencao'))

        res = db_query("SELECT id_ativo FROM manutencoes WHERE id = :id", {'id': id})
        if not res:
            flash('Registro de manutenção não encontrado.', 'error')
            return redirect(url_for('main.manutencao'))
            
        id_ativo = res[0]['id_ativo']
        
        db_execute("DELETE FROM manutencoes WHERE id = :id", {'id': id})
        
        db_execute("""
            INSERT INTO historico (id_ativo, evento, detalhes)
            VALUES (:id_ativo, 'Movimentação', 'Registro de manutenção excluído pelo administrador.')
        """, {'id_ativo': id_ativo})

        flash('Registro de manutenção excluído com sucesso!', 'success')
    except Exception as e:
        flash(f"Erro ao excluir manutenção: {e}", 'error')
        
    return redirect(url_for('main.manutencao'))


@main_bp.route('/relatorios')
@login_required
def relatorios():
    where_sql, params, filtros_selecionados = get_filter_clauses_and_params()
    todas_categorias = db_query("SELECT id, nome FROM categorias ORDER BY nome")
    todos_modelos = db_query("SELECT id, nome, categoria_id FROM modelos ORDER BY nome")
    ativos = db_query(f"""
        SELECT a.id_ativo, c.nome as categoria, m.nome as modelo, a.numero_serie,
               a.status, a.usuario_responsavel, a.localizacao
        FROM ativos a
        JOIN modelos m ON a.modelo_id = m.id
        JOIN categorias c ON a.categoria_id = c.id
        {where_sql}
        ORDER BY c.nome, m.nome
    """, params)

    consumiveis = db_query("SELECT * FROM consumiveis ORDER BY nome")

    return render_template(
        "relatorio_geral.html",
        title="Relatório Geral de TI",
        ativos=ativos,
        total=len(ativos),
        consumiveis=consumiveis,
        total_consumiveis=len(consumiveis),
        data_geracao=datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
        todas_categorias=todas_categorias,
        todos_modelos=todos_modelos,
        filtros_selecionados=filtros_selecionados,
    )


def _enviar_email_assinatura(email_usuario, usuario, id_ativo, token):
    import threading
    from flask import current_app
    from app.models import get_setting

    # Lê configurações do banco de dados com fallback para as do config.py
    server_host = get_setting('mail_server') or current_app.config.get('MAIL_SERVER')
    
    port_setting = get_setting('mail_port')
    if port_setting:
        try:
            port = int(port_setting)
        except ValueError:
            port = current_app.config.get('MAIL_PORT') or 587
    else:
        port = current_app.config.get('MAIL_PORT') or 587

    user = get_setting('mail_username') or current_app.config.get('MAIL_USERNAME')
    password = get_setting('mail_password') or current_app.config.get('MAIL_PASSWORD')
    sender = get_setting('mail_default_sender') or current_app.config.get('MAIL_DEFAULT_SENDER') or user
    
    use_tls_setting = get_setting('mail_use_tls')
    if use_tls_setting is not None:
        use_tls = use_tls_setting.lower() in ('true', '1', 'yes')
    else:
        use_tls = current_app.config.get('MAIL_USE_TLS')

    # Validação rápida de credenciais padrão ou ausência de dados
    if not email_usuario or not user or not password or user == 'seu-email@gmail.com' or password == 'sua-senha-de-aplicativo':
        return False

    try:
        link = url_for('main.assinar_termo', token=token, _external=True)
        company_name = get_setting('company_name', 'Almoxarifado Digital')
        
        msg_content = f"""Olá {usuario},

Um Termo de Responsabilidade para o ativo {id_ativo} foi gerado e aguarda sua assinatura digital.

Por favor, acesse o link abaixo para visualizar os termos e realizar a assinatura digital de recebimento:
{link}

Atenciosamente,
Setor de TI / {company_name}
"""
        subject = f"Assinatura Digital de Termo - Ativo {id_ativo}"

        def send_email_thread():
            import smtplib
            from email.mime.text import MIMEText
            from email.header import Header
            try:
                msg = MIMEText(msg_content, 'plain', 'utf-8')
                msg['Subject'] = Header(subject, 'utf-8')
                msg['From'] = sender
                msg['To'] = email_usuario

                with smtplib.SMTP(server_host, port, timeout=10) as server:
                    if use_tls:
                        server.starttls()
                    server.login(user, password)
                    server.sendmail(sender, [email_usuario], msg.as_string())
                print(f"E-mail enviado com sucesso em segundo plano para: {email_usuario}")
            except Exception as thread_error:
                print(f"Erro ao enviar e-mail em segundo plano para {email_usuario}: {thread_error}")

        # Inicia o envio em segundo plano
        t = threading.Thread(target=send_email_thread, daemon=True)
        t.start()
        return True
    except Exception as e:
        print(f"Erro ao iniciar thread de envio de e-mail: {e}")
        return False


@main_bp.route('/ativo/<id_ativo>/distribuir', methods=['POST'])
@login_required
@role_required('admin', 'tecnico')
def distribuir_ativo(id_ativo):
    import secrets
    form_data = request.form
    template_name = form_data.get('template_name', '')
    _validate_template_name(template_name)

    localidade = form_data.get('localidade', '').strip()
    setor = form_data.get('setor', '').strip()
    localizacao_final = f"{localidade} - {setor}" if localidade and setor else localidade or setor
    usuario_resp = form_data.get('usuario_responsavel', '').strip()
    email_usuario = form_data.get('email_usuario', '').strip()
    unidade = form_data.get('unidade', '').strip()
    chamado = form_data.get('numero_chamado', '').strip()

    detalhes = {
        'usuario_responsavel': usuario_resp,
        'localizacao': localizacao_final,
        'destino': setor or localizacao_final,
    }

    # 1. Atualizar o status do ativo
    manager.movimentar(id_ativo, 'Em Uso', chamado, detalhes)

    # 2. Gerar token de assinatura digital
    token = secrets.token_urlsafe(32)

    # 3. Registrar o termo pendente no banco de dados
    try:
        db_execute("""
            INSERT INTO termos (id_ativo, solicitante, usuario, email_usuario, unidade, localidade, setor, chamado, template_name, token, assinado)
            VALUES (:id_ativo, :solicitante, :usuario, :email_usuario, :unidade, :localidade, :setor, :chamado, :template_name, :token, 0)
        """, {
            'id_ativo': id_ativo,
            'solicitante': current_user.username,
            'usuario': usuario_resp,
            'email_usuario': email_usuario or None,
            'unidade': unidade or None,
            'localidade': localidade or None,
            'setor': setor or None,
            'chamado': chamado or None,
            'template_name': template_name,
            'token': token
        })

        # 4. Tentar enviar o e-mail
        email_enviado = _enviar_email_assinatura(email_usuario, usuario_resp, id_ativo, token)
        if email_enviado:
            flash('Ativo distribuído com sucesso! E-mail com link de assinatura enviado para o colaborador.', 'success')
        else:
            flash('Ativo distribuído com sucesso! Link de assinatura digital gerado.', 'success')
    except Exception as e:
        flash(f'Erro ao registrar termo de assinatura: {e}', 'error')

    return redirect(url_for('main.detalhes_ativo', id_ativo=id_ativo))


@main_bp.route('/assinar/<token>', methods=['GET', 'POST'])
def assinar_termo(token):
    term_list = db_query("SELECT * FROM termos WHERE token = :token", {'token': token})
    if not term_list:
        abort(404, description="Link de assinatura inválido ou expirado.")
    term = term_list[0]
    id_ativo = term['id_ativo']

    query_ativo = """
        SELECT a.*, m.nome as modelo, c.nome as categoria
        FROM ativos a
        JOIN modelos m ON a.modelo_id = m.id
        JOIN categorias c ON a.categoria_id = c.id
        WHERE a.id_ativo = :id_ativo
    """
    ativo_list = db_query(query_ativo, {'id_ativo': id_ativo})
    if not ativo_list:
        abort(404, description="Ativo associado ao termo não encontrado.")
    ativo = ativo_list[0]

    if request.method == 'POST':
        if term['assinado']:
            flash('Este termo já foi assinado.', 'info')
            return redirect(url_for('main.assinar_termo', token=token))

        ip = request.remote_addr
        now_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        try:
            db_execute("""
                UPDATE termos 
                SET assinado = 1, data_assinatura = :now, ip_assinatura = :ip
                WHERE token = :token
            """, {
                'now': now_str,
                'ip': ip,
                'token': token
            })

            detalhes_log = f"Termo de Responsabilidade assinado digitalmente por {term['usuario']} (IP: {ip})."
            db_execute("""
                INSERT INTO historico (id_ativo, evento, detalhes)
                VALUES (:id_ativo, 'Assinatura Digital', :detalhes)
            """, {'id_ativo': id_ativo, 'detalhes': detalhes_log})

            flash('Termo assinado digitalmente com sucesso!', 'success')
            return redirect(url_for('main.assinar_termo', token=token))
        except Exception as e:
            flash(f"Erro ao assinar termo: {e}", 'error')

    return render_template('assinar_termo.html', term=term, ativo=ativo)


@main_bp.route('/ativo/<id_ativo>/movimentar', methods=['POST'])
@login_required
@role_required('admin', 'tecnico')
def movimentar_ativo(id_ativo):
    novo_status = _normalize_status(request.form.get('novo_status', '').strip())
    if novo_status not in current_app.config['MOVEMENT_STATUSES']:
        flash('Status inválido para movimentação.', 'error')
        return redirect(url_for('main.detalhes_ativo', id_ativo=id_ativo))

    # Impede mover o ativo para fora de manutenção por esta rota simples
    ativo_list = db_query("SELECT status FROM ativos WHERE id_ativo = :id_ativo", {'id_ativo': id_ativo})
    if ativo_list and ativo_list[0]['status'] == 'Em Manutenção' and novo_status != 'Em Manutenção':
        flash('Para retirar o ativo de manutenção, utilize a opção "Concluir Manutenção" para registrar os custos e detalhes do reparo.', 'error')
        return redirect(url_for('main.detalhes_ativo', id_ativo=id_ativo))

    chamado = request.form.get('numero_chamado', '').strip()
    manager.movimentar(id_ativo, novo_status, chamado)

    if novo_status == 'Em Manutenção':
        # Evitar duplicar registros ativos de manutenção
        existing = db_query("SELECT id FROM manutencoes WHERE id_ativo = :id_ativo AND status = 'Em Manutenção'", {'id_ativo': id_ativo})
        if not existing:
            db_execute("""
                INSERT INTO manutencoes (id_ativo, numero_chamado, data_inicio, status)
                VALUES (:id_ativo, :chamado, :data_inicio, 'Em Manutenção')
            """, {
                'id_ativo': id_ativo,
                'chamado': chamado,
                'data_inicio': datetime.now().strftime('%Y-%m-%d')
            })

    flash(f'Status do ativo alterado para "{novo_status}" com sucesso!', 'success')
    return redirect(url_for('main.detalhes_ativo', id_ativo=id_ativo))


@main_bp.route('/ativo/<id_ativo>/baixar', methods=['POST'])
@login_required
@role_required('admin', 'tecnico')
def baixar_ativo(id_ativo):
    chamado = request.form.get('numero_chamado', '').strip()
    manager.baixar(id_ativo, chamado)
    flash('Ativo baixado (descartado) com sucesso!', 'success')
    return redirect(url_for('main.detalhes_ativo', id_ativo=id_ativo))


@main_bp.route('/gerenciar', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'tecnico')
def gerenciar_tipos():
    if request.method == 'POST':
        form_type = request.form.get('form_type')
        nome = request.form.get('nome', '').strip()
        if not nome:
            flash('Informe um nome válido.', 'error')
            return redirect(url_for('main.gerenciar_tipos'))

        try:
            if form_type == 'categoria':
                db_execute("INSERT INTO categorias (nome) VALUES (:nome)", {'nome': nome})
                flash(f"Categoria '{nome}' adicionada com sucesso!", 'success')
            elif form_type == 'modelo':
                categoria_id = request.form.get('categoria_id', type=int)
                if not categoria_id:
                    raise ValueError('Selecione uma categoria.')
                db_execute(
                    "INSERT INTO modelos (nome, categoria_id) VALUES (:nome, :categoria_id)",
                    {'nome': nome, 'categoria_id': categoria_id},
                )
                flash(f"Modelo '{nome}' adicionado com sucesso!", 'success')
            else:
                flash('Tipo de formulário inválido.', 'error')
        except Exception as e:
            flash(f"Erro ao salvar: {e}", 'error')
        return redirect(url_for('main.gerenciar_tipos'))

    categorias = db_query("SELECT * FROM categorias ORDER BY nome")
    return render_template("gerenciar_tipos.html", title="Gerenciar Categorias e Modelos", categorias=categorias)


@main_bp.route('/importar', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'tecnico')
def importar_ativos():
    if request.method == 'POST':
        file = request.files.get('file')
        if not file or not file.filename:
            flash('Nenhum arquivo selecionado', 'error')
            return redirect(request.url)

        try:
            df = _read_import_file(file)
            column_mapping = {
                'ativo': 'categoria',
                'fabricante': 'marca',
                'nº serie': 'numero_serie',
                'n° serie': 'numero_serie',
                'n serie': 'numero_serie',
                'n serie': 'numero_serie',
                'numero serie': 'numero_serie',
                'numero_serie': 'numero_serie',
                'modelo': 'modelo',
                'quantidade': 'quantidade',
            }
            df.columns = [str(col).lower().strip().replace('.', '') for col in df.columns]
            df.rename(columns=column_mapping, inplace=True)

            required_columns = ['categoria', 'marca', 'modelo']
            if not all(col in df.columns for col in required_columns):
                flash(f'O arquivo deve conter as colunas obrigatórias: {required_columns}', 'error')
                return redirect(request.url)

            sucesso = 0
            erros = []
            asset_manager = AssetManager()

            for index, row in df.iterrows():
                try:
                    quantidade = int(row.get('quantidade', 1) or 1)
                    if quantidade < 1:
                        quantidade = 1
                except (TypeError, ValueError):
                    quantidade = 1

                base_numero_serie = str(row.get('numero_serie', '')).strip()
                categoria_nome = str(row['categoria']).strip()
                if not categoria_nome:
                    continue

                try:
                    cat_query = db_query("SELECT id FROM categorias WHERE nome = :nome", {'nome': categoria_nome})
                    if not cat_query:
                        db_execute("INSERT INTO categorias (nome) VALUES (:nome)", {'nome': categoria_nome})
                        cat_query = db_query("SELECT id FROM categorias WHERE nome = :nome", {'nome': categoria_nome})
                    cat_id = cat_query[0]['id']

                    modelo_nome = str(row.get('modelo', '')).strip() or categoria_nome
                    mod_query = db_query(
                        "SELECT id FROM modelos WHERE nome = :nome AND categoria_id = :cat_id",
                        {'nome': modelo_nome, 'cat_id': cat_id},
                    )
                    if not mod_query:
                        db_execute(
                            "INSERT INTO modelos (nome, categoria_id) VALUES (:nome, :cat_id)",
                            {'nome': modelo_nome, 'cat_id': cat_id},
                        )
                        mod_query = db_query(
                            "SELECT id FROM modelos WHERE nome = :nome AND categoria_id = :cat_id",
                            {'nome': modelo_nome, 'cat_id': cat_id},
                        )
                    mod_id = mod_query[0]['id']

                    for i in range(quantidade):
                        current_numero_serie = base_numero_serie
                        if base_numero_serie and i > 0:
                            match = re.search(r'(\d+)$', base_numero_serie)
                            if match:
                                num_part = int(match.group(1)) + i
                                prefix = base_numero_serie[:match.start(1)]
                                current_numero_serie = f"{prefix}{str(num_part).zfill(len(match.group(1)))}"
                            else:
                                current_numero_serie = f"{base_numero_serie}-{i + 1}"

                        form_data = {
                            'id_ativo': '',
                            'numero_serie': current_numero_serie,
                            'marca': str(row.get('marca', 'N/A')).strip() or 'N/A',
                            'modelo': mod_id,
                            'categoria': cat_id,
                            'data_aquisicao': None,
                            'destino': 'Estoque TI',
                            'nota_fiscal': row.get('nota_fiscal', ''),
                            'fornecedor': row.get('fornecedor', ''),
                            'cpu': row.get('cpu', ''),
                            'ram_gb': row.get('ram_gb', None),
                            'armazenamento_gb': row.get('armazenamento_gb', None),
                            'sistema_operacional': row.get('sistema_operacional', ''),
                        }
                        asset_manager.registrar_novo_ativo(form_data)
                        sucesso += 1
                except Exception as e:
                    erros.append(f"Linha {index + 2}: {e}")

            if not erros:
                flash(f'Importação concluída! {sucesso} ativos importados com sucesso.', 'success')
            else:
                erros_str = " | ".join(erros[:3])
                flash(f'{sucesso} ativos importados. Falhas: {len(erros)}. Detalhes: {erros_str}', 'error')
        except Exception as e:
            flash(f'Ocorreu um erro fatal ao processar o arquivo: {e}', 'error')

        return redirect(url_for('main.importar_ativos'))

    return render_template('importar.html', title="Importar Ativos")


@main_bp.route('/gerar-termo')
@login_required
@role_required('admin', 'tecnico')
def gerador_termo():
    return render_template('gerador_termo.html', title="Gerador de Termo")


@main_bp.route('/ativo/<id_ativo>/edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'tecnico')
def edit_ativo(id_ativo):
    if request.method == 'POST':
        try:
            novo_id_ativo = manager.atualizar_ativo(id_ativo, request.form)
            flash('Ativo atualizado com sucesso!', 'success')
            return redirect(url_for('main.detalhes_ativo', id_ativo=novo_id_ativo))
        except ValueError as e:
            flash(f'Erro ao atualizar: {e}', 'error')
            return redirect(url_for('main.edit_ativo', id_ativo=id_ativo))

    ativo_list = db_query("SELECT * FROM ativos WHERE id_ativo = :id_ativo", {'id_ativo': id_ativo})
    if not ativo_list:
        flash('Ativo não encontrado.', 'error')
        return redirect(url_for('main.index'))

    categorias = db_query("SELECT * FROM categorias ORDER BY nome")
    modelos = db_query("SELECT * FROM modelos ORDER BY nome")
    return render_template(
        'edit_ativo.html',
        title="Editar Ativo",
        ativo=ativo_list[0],
        categorias=categorias,
        modelos=modelos,
    )


@main_bp.route('/bulk-edit', methods=['GET', 'POST'])
@login_required
@role_required('admin', 'tecnico')
def bulk_edit_ativos():
    if request.method == 'POST':
        try:
            manager.atualizar_lote_ativos(request.form)
            flash('Ativos atualizados com sucesso!', 'success')
            return redirect(url_for('main.bulk_edit_ativos'))
        except ValueError as e:
            flash(f'Erro ao salvar: {e}', 'error')

    ativos_para_editar = db_query("""
        SELECT a.id, a.id_ativo, a.numero_serie, c.nome as categoria, m.nome as modelo
        FROM ativos a
        JOIN modelos m ON a.modelo_id = m.id
        JOIN categorias c ON a.categoria_id = c.id
        WHERE a.numero_serie LIKE 'PROV-%'
        ORDER BY a.id
    """)
    return render_template('bulk_edit.html', title="Editar Ativos em Lote", ativos=ativos_para_editar)


@main_bp.route('/ativo/<id_ativo>/gerar_documento/<template_name>')
@login_required
def gerar_documento(id_ativo, template_name):
    template_path = _validate_template_name(template_name)
    try:
        chamado = request.args.get('chamado', '____________________')
        unidade_form = request.args.get('unidade', '____________________')
        email_usuario_form = request.args.get('email_usuario', '____________________')
        localidade_form = request.args.get('localidade', '____________________')
        setor_form = request.args.get('setor', '____________________')

        query = """
            SELECT a.*, m.nome as modelo, c.nome as categoria
            FROM ativos a
            JOIN modelos m ON a.modelo_id = m.id
            JOIN categorias c ON a.categoria_id = c.id
            WHERE a.id_ativo = :id_ativo
        """
        ativo_info_list = db_query(query, {'id_ativo': id_ativo})
        if not ativo_info_list:
            flash('Ativo não encontrado.', 'error')
            return redirect(url_for('main.index'))
        ativo_info = ativo_info_list[0]

        context = {
            '{{solicitante}}': ativo_info.get('usuario_responsavel', '____________________'),
            '{{usuario}}': email_usuario_form,
            '{{matricula}}': '____________________',
            '{{setor}}': setor_form,
            '{{unidade}}': unidade_form,
            '{{localidade}}': localidade_form,
            '{{patrimonio}}': ativo_info.get('id_ativo', ''),
            '{{categoria}}': ativo_info.get('categoria', ''),
            '{{fabricante}}': ativo_info.get('marca', ''),
            '{{modelo}}': ativo_info.get('modelo', ''),
            '{{serie}}': ativo_info.get('numero_serie', ''),
            '{{chamado}}': chamado,
            '{{data_hoje}}': datetime.now().strftime('%d/%m/%Y'),
        }

        document = Document(str(template_path))
        for paragraph in document.paragraphs:
            for key, value in context.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        run.text = run.text.replace(key, str(value))

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in context.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace(key, str(value))

        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        download_base = template_name.replace(' - modelo.docx', '')
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=f"{download_base}-{id_ativo}.docx",
        )
    except Exception as e:
        flash(f'Ocorreu um erro ao gerar o documento: {e}', 'error')
        return redirect(url_for('main.detalhes_ativo', id_ativo=id_ativo))


@main_bp.route('/almoxarifado')
@login_required
def almoxarifado():
    try:
        q = request.args.get('q', '').strip()
        query = "SELECT * FROM consumiveis"
        params = {}
        if q:
            query += " WHERE nome LIKE :search_term"
            params['search_term'] = f"%{q}%"
        query += " ORDER BY nome"
        
        itens = db_query(query, params)
        total_itens = len(itens)
        itens_alerta = sum(1 for item in itens if item['quantidade'] <= item['estoque_minimo'])
        
        return render_template(
            'almoxarifado.html',
            title="Almoxarifado Geral",
            itens=itens,
            total_itens=total_itens,
            itens_alerta=itens_alerta,
            q=q
        )
    except Exception as e:
        flash(f"Erro ao carregar almoxarifado: {e}", "error")
        return redirect(url_for('main.index'))


@main_bp.route('/almoxarifado/novo', methods=['GET', 'POST'])
@login_required
def almoxarifado_novo():
    if request.method == 'POST':
        try:
            nome = request.form.get('nome', '').strip()
            quantidade = int(request.form.get('quantidade', 0))
            unidade_medida = request.form.get('unidade_medida', 'unidade').strip()
            estoque_minimo = int(request.form.get('estoque_minimo', 0))
            localizacao = request.form.get('localizacao', '').strip()
            fornecedor = request.form.get('fornecedor', '').strip()
            observacoes = request.form.get('observacoes', '').strip()
            
            if not nome:
                flash("O nome do consumível é obrigatório.", "error")
                return redirect(url_for('main.almoxarifado_novo'))
                
            if quantidade < 0 or estoque_minimo < 0:
                flash("A quantidade e o estoque mínimo não podem ser negativos.", "error")
                return redirect(url_for('main.almoxarifado_novo'))
                
            db_execute("""
                INSERT INTO consumiveis (nome, quantidade, unidade_medida, estoque_minimo, localizacao, fornecedor, observacoes)
                VALUES (:nome, :quantidade, :unidade_medida, :estoque_minimo, :localizacao, :fornecedor, :observacoes)
            """, {
                'nome': nome,
                'quantidade': quantidade,
                'unidade_medida': unidade_medida,
                'estoque_minimo': estoque_minimo,
                'localizacao': localizacao,
                'fornecedor': fornecedor,
                'observacoes': observacoes
            })
            
            res = db_query("SELECT id FROM consumiveis WHERE nome = :nome", {'nome': nome})
            if res:
                consumivel_id = res[0]['id']
                db_execute("""
                    INSERT INTO consumiveis_historico (consumivel_id, tipo_movimentacao, quantidade, usuario, detalhes)
                    VALUES (:consumivel_id, 'Entrada', :quantidade, :usuario, 'Cadastro inicial do item')
                """, {
                    'consumivel_id': consumivel_id,
                    'quantidade': quantidade,
                    'usuario': current_user.username
                })
            
            flash("Item cadastrado com sucesso no almoxarifado!", "success")
            return redirect(url_for('main.almoxarifado'))
        except Exception as e:
            flash(f"Erro ao cadastrar consumível: {e}", "error")
            return redirect(url_for('main.almoxarifado_novo'))
            
    return render_template('almoxarifado_novo.html', title="Novo Item de Consumo")


@main_bp.route('/almoxarifado/movimentar/<int:item_id>', methods=['POST'])
@login_required
def almoxarifado_movimentar(item_id):
    try:
        tipo_movimentacao = request.form.get('tipo_movimentacao', '').strip()
        quantidade = int(request.form.get('quantidade', 0))
        numero_chamado = request.form.get('numero_chamado', '').strip() or None
        detalhes = request.form.get('detalhes', '').strip() or None
        
        if quantidade <= 0:
            flash("A quantidade para movimentação deve ser maior que zero.", "error")
            return redirect(url_for('main.almoxarifado'))
            
        res = db_query("SELECT * FROM consumiveis WHERE id = :id", {'id': item_id})
        if not res:
            flash("Item não encontrado.", "error")
            return redirect(url_for('main.almoxarifado'))
            
        item = res[0]
        q_atual = item['quantidade']
        
        if tipo_movimentacao == 'Entrada':
            nova_q = q_atual + quantidade
        elif tipo_movimentacao == 'Saída':
            if q_atual < quantidade:
                flash(f"Estoque insuficiente para a retirada. Disponível: {q_atual} {item['unidade_medida']}.", "error")
                return redirect(url_for('main.almoxarifado'))
            nova_q = q_atual - quantidade
        elif tipo_movimentacao == 'Ajuste':
            nova_q = quantidade
        else:
            flash("Tipo de movimentação inválido.", "error")
            return redirect(url_for('main.almoxarifado'))
            
        db_execute("UPDATE consumiveis SET quantidade = :quantidade WHERE id = :id", {
            'quantidade': nova_q,
            'id': item_id
        })
        
        db_execute("""
            INSERT INTO consumiveis_historico (consumivel_id, tipo_movimentacao, quantidade, usuario, numero_chamado, detalhes)
            VALUES (:consumivel_id, :tipo_movimentacao, :quantidade, :usuario, :numero_chamado, :detalhes)
        """, {
            'consumivel_id': item_id,
            'tipo_movimentacao': tipo_movimentacao,
            'quantidade': quantidade,
            'usuario': current_user.username,
            'numero_chamado': numero_chamado,
            'detalhes': detalhes
        })
        
        flash(f"Movimentação de {tipo_movimentacao} realizada com sucesso!", "success")
    except Exception as e:
        flash(f"Erro ao processar movimentação: {e}", "error")
        
    return redirect(url_for('main.almoxarifado'))


@main_bp.route('/almoxarifado/historico')
@login_required
def almoxarifado_historico():
    try:
        historico = db_query("""
            SELECT h.*, c.nome as consumivel_nome, c.unidade_medida
            FROM consumiveis_historico h
            JOIN consumiveis c ON h.consumivel_id = c.id
            ORDER BY h.data_movimentacao DESC
        """)
        return render_template(
            'almoxarifado_historico.html',
            title="Histórico do Almoxarifado",
            historico=historico
        )
    except Exception as e:
        flash(f"Erro ao carregar histórico: {e}", "error")
        return redirect(url_for('main.almoxarifado'))


@main_bp.route('/almoxarifado/importar', methods=['GET', 'POST'])
@login_required
def almoxarifado_importar():
    if request.method == 'POST':
        file = request.files.get('file')
        if not file or not file.filename:
            flash('Nenhum arquivo selecionado', 'error')
            return redirect(request.url)

        try:
            df = _read_import_file(file)
            column_mapping = {
                'nome': 'nome',
                'quantidade': 'quantidade',
                'unidade_medida': 'unidade_medida',
                'unidade': 'unidade_medida',
                'estoque_minimo': 'estoque_minimo',
                'estoque minimo': 'estoque_minimo',
                'localizacao': 'localizacao',
                'localização': 'localizacao',
                'fornecedor': 'fornecedor',
                'observacoes': 'observacoes',
                'observação': 'observacoes',
                'observações': 'observacoes',
            }
            df.columns = [str(col).lower().strip().replace('.', '') for col in df.columns]
            df.rename(columns=column_mapping, inplace=True)

            if 'nome' not in df.columns:
                flash("O arquivo deve conter pelo menos a coluna 'nome'.", 'error')
                return redirect(request.url)

            sucesso = 0
            erros = []

            for index, row in df.iterrows():
                try:
                    nome = str(row['nome']).strip()
                    if not nome:
                        continue

                    try:
                        quantidade = int(row.get('quantidade', 0) or 0)
                        if quantidade < 0:
                            quantidade = 0
                    except (TypeError, ValueError):
                        quantidade = 0

                    try:
                        estoque_minimo = int(row.get('estoque_minimo', 0) or 0)
                        if estoque_minimo < 0:
                            estoque_minimo = 0
                    except (TypeError, ValueError):
                        estoque_minimo = 0

                    unidade_medida = str(row.get('unidade_medida', 'unidade')).strip() or 'unidade'
                    localizacao = str(row.get('localizacao', '')).strip() or None
                    fornecedor = str(row.get('fornecedor', '')).strip() or None
                    observacoes = str(row.get('observacoes', '')).strip() or None

                    res = db_query("SELECT id, quantidade FROM consumiveis WHERE nome = :nome", {'nome': nome})
                    if res:
                        item_id = res[0]['id']
                        q_antiga = res[0]['quantidade']
                        nova_q = q_antiga + quantidade

                        db_execute("""
                            UPDATE consumiveis 
                            SET quantidade = :quantidade,
                                localizacao = COALESCE(:localizacao, localizacao),
                                fornecedor = COALESCE(:fornecedor, fornecedor),
                                observacoes = COALESCE(:observacoes, observacoes)
                            WHERE id = :id
                        """, {
                            'quantidade': nova_q,
                            'localizacao': localizacao,
                            'fornecedor': fornecedor,
                            'observacoes': observacoes,
                            'id': item_id
                        })

                        if quantidade > 0:
                            db_execute("""
                                INSERT INTO consumiveis_historico (consumivel_id, tipo_movimentacao, quantidade, usuario, detalhes)
                                VALUES (:consumivel_id, 'Entrada', :quantidade, :usuario, 'Aumento de estoque via importação')
                            """, {
                                'consumivel_id': item_id,
                                'quantidade': quantidade,
                                'usuario': current_user.username
                            })
                    else:
                        db_execute("""
                            INSERT INTO consumiveis (nome, quantidade, unidade_medida, estoque_minimo, localizacao, fornecedor, observacoes)
                            VALUES (:nome, :quantidade, :unidade_medida, :estoque_minimo, :localizacao, :fornecedor, :observacoes)
                        """, {
                            'nome': nome,
                            'quantidade': quantidade,
                            'unidade_medida': unidade_medida,
                            'estoque_minimo': estoque_minimo,
                            'localizacao': localizacao,
                            'fornecedor': fornecedor,
                            'observacoes': observacoes
                        })

                        new_res = db_query("SELECT id FROM consumiveis WHERE nome = :nome", {'nome': nome})
                        if new_res:
                            item_id = new_res[0]['id']
                            db_execute("""
                                INSERT INTO consumiveis_historico (consumivel_id, tipo_movimentacao, quantidade, usuario, detalhes)
                                VALUES (:consumivel_id, 'Entrada', :quantidade, :usuario, 'Importação inicial do item')
                            """, {
                                'consumivel_id': item_id,
                                'quantidade': quantidade,
                                'usuario': current_user.username
                            })

                    sucesso += 1
                except Exception as e:
                    erros.append(f"Linha {index + 2}: {e}")

            if not erros:
                flash(f'Importação concluída! {sucesso} consumíveis importados com sucesso.', 'success')
            else:
                erros_str = " | ".join(erros[:3])
                flash(f'{sucesso} consumíveis importados. Falhas: {len(erros)}. Detalhes: {erros_str}', 'error')

            return redirect(url_for('main.almoxarifado'))
        except Exception as e:
            flash(f"Erro ao processar arquivo: {e}", 'error')
            return redirect(request.url)

    return render_template('almoxarifado_importar.html', title="Importar Consumíveis")


@main_bp.route('/almoxarifado/item/<int:item_id>')
@login_required
def almoxarifado_item_detalhes(item_id):
    try:
        res = db_query("SELECT * FROM consumiveis WHERE id = :id", {'id': item_id})
        if not res:
            flash("Item não encontrado no almoxarifado.", "error")
            return redirect(url_for('main.almoxarifado'))
        
        item = res[0]
        
        historico = db_query("""
            SELECT * FROM consumiveis_historico 
            WHERE consumivel_id = :id 
            ORDER BY data_movimentacao DESC
        """, {'id': item_id})
        
        total_operacoes = len(historico)
        
        return render_template(
            'almoxarifado_detalhes.html',
            title=f"Detalhes: {item['nome']}",
            item=item,
            historico=historico,
            total_operacoes=total_operacoes
        )
    except Exception as e:
        flash(f"Erro ao carregar detalhes: {e}", "error")
        return redirect(url_for('main.almoxarifado'))


@main_bp.route('/almoxarifado/item/<int:item_id>/editar', methods=['GET', 'POST'])
@login_required
def almoxarifado_item_editar(item_id):
    res = db_query("SELECT * FROM consumiveis WHERE id = :id", {'id': item_id})
    if not res:
        flash("Item não encontrado no almoxarifado.", "error")
        return redirect(url_for('main.almoxarifado'))
    
    item = res[0]
    
    if request.method == 'POST':
        try:
            nome = request.form.get('nome', '').strip()
            unidade_medida = request.form.get('unidade_medida', 'unidade').strip()
            estoque_minimo = int(request.form.get('estoque_minimo', 0))
            localizacao = request.form.get('localizacao', '').strip() or None
            fornecedor = request.form.get('fornecedor', '').strip() or None
            observacoes = request.form.get('observacoes', '').strip() or None
            
            if not nome:
                flash("O nome do consumível é obrigatório.", "error")
                return redirect(url_for('main.almoxarifado_item_editar', item_id=item_id))
                
            if estoque_minimo < 0:
                flash("O estoque mínimo não pode ser negativo.", "error")
                return redirect(url_for('main.almoxarifado_item_editar', item_id=item_id))

            if nome.lower() != item['nome'].lower():
                duplicado = db_query("SELECT id FROM consumiveis WHERE LOWER(nome) = LOWER(:nome)", {'nome': nome})
                if duplicado:
                    flash(f"Já existe outro item cadastrado com o nome '{nome}'.", "error")
                    return redirect(url_for('main.almoxarifado_item_editar', item_id=item_id))

            db_execute("""
                UPDATE consumiveis
                SET nome = :nome,
                    unidade_medida = :unidade_medida,
                    estoque_minimo = :estoque_minimo,
                    localizacao = :localizacao,
                    fornecedor = :fornecedor,
                    observacoes = :observacoes
                WHERE id = :id
            """, {
                'nome': nome,
                'unidade_medida': unidade_medida,
                'estoque_minimo': estoque_minimo,
                'localizacao': localizacao,
                'fornecedor': fornecedor,
                'observacoes': observacoes,
                'id': item_id
            })

            db_execute("""
                INSERT INTO consumiveis_historico (consumivel_id, tipo_movimentacao, quantidade, usuario, detalhes)
                VALUES (:consumivel_id, 'Ajuste', 0, :usuario, 'Dados do cadastro atualizados')
            """, {
                'consumivel_id': item_id,
                'usuario': current_user.username
            })
            
            flash("Item do almoxarifado atualizado com sucesso!", "success")
            return redirect(url_for('main.almoxarifado_item_detalhes', item_id=item_id))
        except Exception as e:
            flash(f"Erro ao atualizar item: {e}", "error")
            return redirect(url_for('main.almoxarifado_item_editar', item_id=item_id))
            
    return render_template('almoxarifado_editar.html', title=f"Editar: {item['nome']}", item=item)


@main_bp.route('/almoxarifado/item/<int:item_id>/excluir', methods=['POST'])
@login_required
def almoxarifado_item_excluir(item_id):
    try:
        res = db_query("SELECT * FROM consumiveis WHERE id = :id", {'id': item_id})
        if not res:
            flash("Item não encontrado.", "error")
            return redirect(url_for('main.almoxarifado'))
            
        item = res[0]
        db_execute("DELETE FROM consumiveis_historico WHERE consumivel_id = :id", {'id': item_id})
        db_execute("DELETE FROM consumiveis WHERE id = :id", {'id': item_id})
        
        flash(f"Item '{item['nome']}' e todas as suas movimentações foram excluídos.", "success")
    except Exception as e:
        flash(f"Erro ao excluir item do almoxarifado: {e}", "error")
        
    return redirect(url_for('main.almoxarifado'))


@main_bp.route('/ativo/etiquetas-lote', methods=['GET', 'POST'])
@login_required
def etiquetas_lote():
    if request.method == 'POST':
        # Recebe os IDs dos ativos selecionados (id_ativo)
        selected_ids = request.form.getlist('ativos_selecionados')
        if not selected_ids:
            flash('Por favor, selecione pelo menos um ativo para gerar etiquetas.', 'error')
            return redirect(url_for('main.etiquetas_lote'))
            
        # Busca detalhes dos ativos selecionados no banco
        placeholders = ', '.join([f":id_{i}" for i in range(len(selected_ids))])
        params = {f"id_{i}": val for i, val in enumerate(selected_ids)}
        
        query = f"""
            SELECT a.id_ativo, a.numero_serie, a.marca, m.nome as modelo, c.nome as categoria
            FROM ativos a
            JOIN modelos m ON a.modelo_id = m.id
            JOIN categorias c ON a.categoria_id = c.id
            WHERE a.id_ativo IN ({placeholders})
            ORDER BY a.id_ativo DESC
        """
        ativos_selecionados = db_query(query, params)
        
        return render_template(
            'etiquetas_impressao.html',
            title="Impressão de Etiquetas em Lote",
            ativos=ativos_selecionados
        )
        
    # GET request: exibe listagem com filtros
    where_sql, params, filtros_selecionados = get_filter_clauses_and_params()
    todas_categorias = db_query("SELECT id, nome FROM categorias ORDER BY nome")
    todos_modelos = db_query("SELECT id, nome, categoria_id FROM modelos ORDER BY nome")
    
    # Lista ativos
    ativos = db_query(f"""
        SELECT a.id_ativo, c.nome as categoria, m.nome as modelo, a.numero_serie, a.status, a.usuario_responsavel
        FROM ativos a
        JOIN modelos m ON a.modelo_id = m.id
        JOIN categorias c ON a.categoria_id = c.id
        {where_sql}
        ORDER BY a.id_ativo DESC
    """, params)
    
    return render_template(
        'etiquetas_lote.html',
        title="Gerador de Etiquetas em Lote",
        ativos=ativos,
        todas_categorias=todas_categorias,
        todos_modelos=todos_modelos,
        filtros_selecionados=filtros_selecionados
    )

